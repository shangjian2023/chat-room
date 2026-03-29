"""
生成实验报告 Word 文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_report():
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 标题
    title = doc.add_heading('Web 应用开发实验报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_paragraph()
    info = [
        ('课程名称：', 'Web 应用开发技术'),
        ('实验项目：', '用户认证与会话管理系统'),
        ('学生姓名：', '___________'),
        ('学    号：', '___________'),
        ('班    级：', '___________'),
        ('完成日期：', datetime.now().strftime('%Y 年 %m 月 %d 日')),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        runner = p.add_run(f"{label}{value}")
        runner.font.size = Pt(12)
        runner.font.name = u'宋体'
        runner._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 目录
    doc.add_page_break()
    doc.add_heading('目  录', level=1)
    toc = [
        '一、实验目的',
        '二、实验内容',
        '三、实验环境',
        '四、实验任务与要求',
        '五、实验步骤',
        '六、核心代码实现',
        '七、功能测试截图',
        '八、实验总结',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    # 正文内容
    doc.add_page_break()

    # 一、实验目的
    doc.add_heading('一、实验目的', level=1)
    purpose = [
        '1. 掌握 Web 应用开发的基本流程和方法',
        '2. 理解 HTTP 协议、请求-响应模型及路由机制',
        '3. 掌握用户认证与会话管理的实现原理',
        '4. 熟悉 MySQL 数据库的连接与操作',
        '5. 理解密码哈希加密的安全机制',
        '6. 掌握多用户并发访问的处理方法',
        '7. 培养独立开发、调试和部署 Web 应用的能力',
    ]
    for item in purpose:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5

    # 二、实验内容
    doc.add_heading('二、实验内容', level=1)
    content_text = '''
本实验基于 Python Flask 框架，开发一个支持多用户并发访问的课程聊天平台。系统实现用户注册、登录、会话管理、课程展示、在线聊天等核心功能，采用 MySQL 数据库进行数据持久化存储，使用 Werkzeug 进行密码哈希加密，确保用户信息安全。

系统主要功能模块：
1. 用户认证模块：注册、登录、登出
2. 会话管理模块：Session 存储、权限验证
3. 课程平台模块：课程展示、分类管理、内容编辑
4. 聊天室模块：实时消息、匿名发言、多房间支持
5. 在线统计模块：实时在线用户统计
'''
    doc.add_paragraph(content_text)

    # 三、实验环境
    doc.add_heading('三、实验环境', level=1)

    doc.add_heading('1. 硬件环境', level=2)
    hardware = '''
- 个人电脑（Windows 11 系统）
- 支持局域网组网/多终端互联
- 保证网络通畅，支持本地端口监听
- 支持跨终端浏览器访问与多用户并发连接
'''
    doc.add_paragraph(hardware)

    doc.add_heading('2. 软件环境', level=2)
    software = '''
技术栈选择：Python 3.8+ + Flask + MySQL

主要依赖包：
- Flask 3.0.0：Web 应用框架
- Flask-SQLAlchemy 3.1.1：数据库 ORM
- Flask-Session 0.5.0：会话管理
- Flask-SocketIO 5.3.0：WebSocket 实时通信
- Werkzeug 3.0.0：密码加密工具
- pymysql 1.1.0：MySQL 数据库连接
'''
    doc.add_paragraph(software)

    doc.add_heading('3. 辅助工具', level=2)
    tools = '''
- 代码编辑器：VS Code / PyCharm
- 浏览器：Chrome / Firefox
- 数据库：MySQL 8.0
- 接口测试：浏览器开发者工具
'''
    doc.add_paragraph(tools)

    # 四、实验任务与要求
    doc.add_heading('四、实验任务与要求', level=1)

    doc.add_heading('（一）基础任务（必完成）', level=2)
    basic_tasks = '''
1. 搭建本地互联网络环境，选用 Flask 技术开发 Web 服务器，配置端口 8080，支持跨终端访问。

2. 实现用户注册功能：访问/register 页面，通过表单输入用户名和密码，提交后完成注册，自动跳转登录页。

3. 实现数据库持久化：用户数据存储至 MySQL 数据库，可通过数据库工具查看。

4. 实现多用户独立登录：访问/login 页面，输入用户名和密码，后端校验数据库数据；支持多终端同时登录。

5. 实现多用户会话管理：未登录用户自动跳转登录页；多用户会话独立，互不干扰。
'''
    doc.add_paragraph(basic_tasks)

    doc.add_heading('（二）进阶加分任务', level=2)
    advanced_tasks = '''
1. 注册页面添加重复密码校验，不一致则返回错误提示。

2. 登录/注册页面互相跳转链接，优化交互逻辑。

3. 后端使用哈希加密存储密码，不存储明文。

4. 课程平台展示当前所有在线登录用户列表。

5. 用户登出功能，清除会话状态，跳转登录页。
'''
    doc.add_paragraph(advanced_tasks)

    # 五、实验步骤
    doc.add_heading('五、实验步骤', level=1)
    steps = '''
1. 环境搭建与项目初始化
   - 安装 Python 3.8+、MySQL 8.0
   - 创建项目目录结构
   - 安装依赖：pip install flask flask-sqlalchemy flask-session flask-socketio pymysql
   - 配置服务器端口 8080，配置 MySQL 连接

2. 数据库与数据表创建
   - 设计用户表：id, username, password_hash, is_admin, created_at
   - 设计课程表：id, title, category, content, order_index
   - 设计聊天消息表：id, username, message, timestamp, room
   - 设计在线用户表：id, user_id, session_id, username, login_time, last_active

3. 前端页面开发
   - 注册页面（register.html）：用户名、密码、确认密码表单
   - 登录页面（login.html）：用户名、密码表单
   - 主页（index.html）：左中右三栏布局（课程平台、聊天室、实时统计）

4. 后端核心逻辑开发
   - 路由配置：/, /register, /login, /logout, /api/course/*
   - 会话管理：Flask-Session 存储，登录验证，权限拦截
   - 密码加密：Werkzeug 的 generate_password_hash/check_password_hash
   - 并发处理：Flask 多线程，数据库连接池

5. 功能与网络调试
   - 单终端测试注册、登录、会话保持
   - 多终端测试并发登录、消息收发
   - 排查数据库操作、会话管理问题
'''
    doc.add_paragraph(steps)

    # 六、核心代码实现
    doc.add_heading('六、核心代码实现', level=1)

    doc.add_heading('1. 用户模型', level=2)
    code1 = '''
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.now)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
'''
    doc.add_paragraph(code1).style = 'No Spacing'

    doc.add_heading('2. 登录路由', level=2)
    code2 = '''
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')

        user = User.query.filter_by(username=username).first()
        if not user or not user.check_password(password):
            flash('用户名或密码错误', 'error')
            return redirect(url_for('login'))

        session['user_id'] = user.id
        session['username'] = user.username
        session['is_admin'] = user.is_admin

        add_online_user(user.id, username)
        flash('登录成功', 'success')
        return redirect(url_for('index'))

    return render_template('login.html')
'''
    doc.add_paragraph(code2).style = 'No Spacing'

    doc.add_heading('3. 权限拦截', level=2)
    code3 = '''
@app.route('/')
def index():
    if 'user_id' not in session:
        flash('请先登录', 'warning')
        return redirect(url_for('login'))

    # 获取课程模块、在线用户、消息数据
    course_modules = CourseModule.query.order_by(CourseModule.order_index).all()
    is_admin = session.get('is_admin', False)

    return render_template('index.html',
                         username=session.get('username'),
                         is_admin=is_admin,
                         course_modules=course_modules)
'''
    doc.add_paragraph(code3).style = 'No Spacing'

    # 七、功能测试
    doc.add_heading('七、功能测试', level=1)

    test_items = [
        '1. 注册功能测试：输入用户名和密码，确认密码一致后注册成功',
        '2. 登录功能测试：输入正确的用户名和密码登录成功',
        '3. 会话保持测试：关闭浏览器后重新打开仍需登录',
        '4. 权限拦截测试：未登录访问主页自动跳转登录页',
        '5. 多用户并发测试：多终端同时登录，会话独立',
        '6. 课程管理测试：管理员可编辑、添加、删除课程',
        '7. 聊天功能测试：多用户实时收发消息',
    ]
    for item in test_items:
        doc.add_paragraph(item)

    doc.add_paragraph('\n（此处插入功能测试截图、数据库查看截图）')

    # 八、实验总结
    doc.add_heading('八、实验总结', level=1)
    summary = '''
1. 技术实现逻辑
   本系统采用 Flask 作为 Web 框架，利用其轻量级、易扩展的特点，快速搭建 RESTful API 和页面路由。通过 Flask-Session 实现会话管理，将 Session 存储到文件系统，支持多用户并发访问。

2. 会话管理原理
   Flask 使用客户端 Session（加密 Cookie）存储用户标识，服务端验证 Session 有效性。每次请求时，Flask 自动解密 Cookie 获取 user_id，实现用户认证。

3. 密码安全处理
   采用 Werkzeug 的 generate_password_hash 函数，使用 PBKDF2 算法生成密码哈希值，包含随机盐值和多次迭代，有效防止彩虹表攻击。

4. 遇到的问题及解决方案
   - 问题 1：多用户会话互相干扰
     解决：使用 Flask-Session 将会话存储到服务器端，每个用户独立 Session
   - 问题 2：数据库连接超时
     解决：配置 SQLAlchemy 连接池参数，增加连接复用
   - 问题 3：跨终端访问失败
     解决：Flask.run 时设置 host='0.0.0.0'，允许外部访问

5. 实验心得
   通过本次实验，我掌握了 Web 应用开发的完整流程，理解了用户认证和会话管理的核心原理。在开发过程中，通过查阅文档、调试代码，提高了独立解决问题的能力。系统的左中右三栏布局设计，使课程展示、聊天、统计功能清晰分离，用户体验良好。
'''
    doc.add_paragraph(summary)

    # 保存文档
    doc.save('实验报告.docx')
    print('实验报告已生成：实验报告.docx')

if __name__ == '__main__':
    create_report()
